import docx
doc = docx.Document("1.docx")

def find_Tables(Head):
    tables = doc.tables
    for table in tables:
        row = table.rows[0]
        cell = row.cells[0]
        if Head.lower() in cell.text.lower():
            return table

    return False

id_table = find_Tables('Главная таблица')
if not id_table:
    print('Таблица с таким заголовком не найдена')
else:
    print(id_table)
    row = id_table.rows[2]

    for cell in row.cells:
        print(cell.text)


def find_txt_paragraph(txt):
    for paragraph in doc.paragraphs:
        if txt.lower() in paragraph.text.lower():
            return paragraph
    return False

id_paragraph = find_txt_paragraph('иванов')
if not id_paragraph:
    print('Параграф с таким заголовком не найден')
else:
    print(id_paragraph.text)
# requests.get(url, stream=True).headers['Content-length']
